#!/usr/bin/env python3
"""
Extract text from Word documents and export to Excel with structured fields.

Uses table/text parsing + Mercer LLM API to fill empty fields.
"""
from __future__ import annotations

import argparse
import asyncio
import io
import json
import os
import re
import tempfile
import unicodedata
from pathlib import Path

import httpx
import pandas as pd

_script_dir = Path(__file__).resolve().parent
try:
    from dotenv import load_dotenv
    load_dotenv(_script_dir / ".env")
except ImportError:
    pass

WORD_EXTENSIONS = {".doc", ".docx"}

# Mandatory/Required fields from Data Attributes (Job Solutions)
MANDATORY_FIELDS = [
    "Organization Name",
    "Job Code",
    "Position title",
    "Leadership Competency Category",
    "Current Compensation Grade",
    "Comments/Notes",
    "Typically Reports to",
    "Employee ID",
    "Manager ID",
    "Direct report counts",
    "People Management Flag",
    "Job Level",
    "Minimum Experience",
    "Pay grade",
    "Department",
    "Job Title (from Source Job Description)",
    "Job Description",
    "Base Salary",
    "Job Location",
]


def preprocess_text(text: str | None) -> str:
    """Clean and preprocess: remove URLs, HTML, control chars, normalize whitespace."""
    if not text or not isinstance(text, str):
        return ""
    text = re.sub(r"https?://[^\s<>\"']+|www\.[^\s<>\"']+", " ", text, flags=re.I)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)|\[[^\]]*\]\([^)]+\)", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\u200b-\u200f\u2028-\u202f\u2060-\u206f\ufffe\uffff]", "", text)
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[\t\n\r]+", " ", text)
    text = re.sub(r" +", " ", text)
    return text.strip()


def _cell_text(cell) -> str:
    """Recursively get all text from a cell, including nested tables."""
    parts = [p.text for p in cell.paragraphs if p.text.strip()]
    for nested_table in getattr(cell, "tables", []):
        parts.append(_table_text(nested_table))
    return "\n".join(parts)


def _table_text(table) -> str:
    """Recursively get all text from a table, including nested tables."""
    parts = []
    for row in table.rows:
        for cell in row.cells:
            text = _cell_text(cell)
            if text.strip():
                parts.append(text)
    return "\n".join(parts)


def _extract_docx(content: bytes) -> str:
    """Extract all text from .docx: paragraphs, tables (including nested), headers, footers."""
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        parts.append(_table_text(table))
    for section in doc.sections:
        for header in (section.header, section.first_page_header):
            if header:
                for p in header.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
                for t in header.tables:
                    parts.append(_table_text(t))
        for footer in (section.footer, section.first_page_footer):
            if footer:
                for p in footer.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
                for t in footer.tables:
                    parts.append(_table_text(t))
    return "\n".join(parts) if parts else ""


def _extract_with_docx2txt(content: bytes, ext: str) -> str:
    """Extract using docx2txt - parses raw XML, gets tables and all content."""
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
        f.write(content)
        path = f.name
    try:
        import docx2txt
        return docx2txt.process(path) or ""
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def _extract_docx_raw_xml(content: bytes) -> str:
    """Extract all text from docx by parsing XML - catches text in shapes, content controls, etc."""
    import zipfile
    import xml.etree.ElementTree as ET
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    texts = []
    with zipfile.ZipFile(io.BytesIO(content), "r") as z:
        for name in z.namelist():
            if "word/" in name and name.endswith(".xml"):
                try:
                    tree = ET.parse(z.open(name))
                    for el in tree.getroot().iter(f"{{{ns}}}t"):
                        if el.text:
                            texts.append(el.text)
                        if el.tail:
                            texts.append(el.tail)
                except ET.ParseError:
                    pass
    return "".join(texts)


def _normalize_field_match(cell_text: str, field_lower: str) -> bool:
    """Check if cell text matches a field name (exact, prefix, or contains)."""
    ct = cell_text.lower().strip()
    return ct == field_lower or ct.startswith(field_lower + " ") or ct.startswith(field_lower + ":") or field_lower in ct


def _extract_fields_from_tables(content: bytes) -> dict[str, str]:
    """Extract field-value pairs directly from Word tables. Handles:
    - 2-col table: [Field] [Value] per row
    - Multi-col: field in any cell, value in next cell
    """
    from docx import Document
    result = {f: "" for f in MANDATORY_FIELDS}
    field_names_lower = {f.lower().strip(): f for f in MANDATORY_FIELDS}

    def scan_table(table):
        for row in table.rows:
            cells = [_cell_text(c).strip() for c in row.cells]
            for i, cell in enumerate(cells):
                if i + 1 >= len(cells):
                    continue
                cell_lower = cell.lower().strip()
                for fl, field in field_names_lower.items():
                    if result[field]:
                        continue
                    if _normalize_field_match(cell, fl):
                        val = cells[i + 1].strip()
                        if val and val.lower() not in field_names_lower:
                            max_len = 8000 if field == "Job Description" else 1000
                            result[field] = val[:max_len]
                        break

    try:
        doc = Document(io.BytesIO(content))
        for table in doc.tables:
            scan_table(table)
        for section in doc.sections:
            for header in (section.header, section.first_page_header):
                if header:
                    for t in header.tables:
                        scan_table(t)
            for footer in (section.footer, section.first_page_footer):
                if footer:
                    for t in footer.tables:
                        scan_table(t)
    except Exception:
        pass
    return result


def extract_text(content: bytes, filename: str) -> str:
    """Extract text using Python libraries. Tries docx2txt first (comprehensive), fallback to python-docx."""
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".docx":
            raw = _extract_with_docx2txt(content, ".docx")
            docx_text = _extract_docx(content)
            if len(docx_text) > len(raw):
                raw = docx_text
            xml_text = _extract_docx_raw_xml(content)
            if len(xml_text) > len(raw):
                raw = xml_text
        elif ext == ".doc":
            try:
                raw = _extract_with_docx2txt(content, ".doc")
            except Exception:
                raw = "[ERROR] .doc extraction failed. Try converting to .docx."
        else:
            return ""
        return preprocess_text(raw)
    except ImportError as e:
        return f"[ERROR] Install python-docx and docx2txt: pip install python-docx docx2txt"
    except Exception as e:
        return f"[ERROR] Extraction failed: {e}"


def _is_field_label(line: str) -> bool:
    return any(line.lower().startswith(f.lower()) for f in MANDATORY_FIELDS)


def _normalize_key(k: str) -> str:
    return re.sub(r"\s+", " ", k.lower().strip().replace("_", " ").replace("-", " "))


def _match_llm_key(out: dict, field: str) -> str:
    """Find value in LLM output using fuzzy key match."""
    fn = _normalize_key(field)
    for k, v in out.items():
        if _normalize_key(k) == fn and v is not None:
            return str(v).strip()
    return ""


async def _extract_fields_with_llm(text: str, verbose: bool = False) -> dict[str, str]:
    """Use Mercer LLM API to extract field values from document text."""
    url = os.getenv("CORE_API_BASE_URL", "https://stg1.mmc-bedford-int-non-prod-ingress.mgti.mmc.com/coreapi/openai/v1/deployments/mmc-tech-gpt-35-turbo-smart-latest/chat/completions")
    api_key = os.getenv("CORE_API_KEY")
    if not api_key:
        if verbose:
            print("  [LLM] CORE_API_KEY not set, skipping")
        return {}
    if not text or len(text) < 20:
        if verbose:
            print("  [LLM] Text too short, skipping")
        return {}
    field_list = "\n".join(f"- {f}" for f in MANDATORY_FIELDS)
    prompt = f"""Extract these fields from the document. Return ONLY valid JSON. Use "" for missing.

Fields to extract:
{field_list}

Document:
{text[:10000]}

Return JSON with keys exactly: Organization Name, Job Code, Position title, Leadership Competency Category, Current Compensation Grade, Comments/Notes, Typically Reports to, Employee ID, Manager ID, Direct report counts, People Management Flag, Job Level, Minimum Experience, Pay grade, Department, Job Title (from Source Job Description), Job Description, Base Salary, Job Location"""
    try:
        async with httpx.AsyncClient(timeout=90) as client:
            r = await client.post(
                url,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json", "x-api-key": api_key},
                json={
                    "messages": [{"role": "user", "content": prompt}],
                    "model": os.getenv("CORE_API_MODEL", "mmc-tech-gpt-35-turbo-smart-latest"),
                    "temperature": 0,
                    "max_tokens": 4096,
                },
            )
            r.raise_for_status()
            data = r.json()
            content = (data.get("choices") or [{}])[0].get("message", {}).get("content", "") or (data.get("choices") or [{}])[0].get("text", "")
            content = str(content).strip()
            if not content:
                if verbose:
                    print("  [LLM] Empty response")
                return {}
            for start in ("```json", "```"):
                if content.startswith(start):
                    content = content[len(start):].strip()
                if content.endswith("```"):
                    content = content[:-3].strip()
            out = json.loads(content)
            if not isinstance(out, dict):
                return {}
            result = {}
            for f in MANDATORY_FIELDS:
                val = out.get(f) or _match_llm_key(out, f)
                if val:
                    result[f] = val[:8000] if f == "Job Description" else val[:1000]
                else:
                    result[f] = ""
            return result
    except Exception as e:
        if verbose:
            print(f"  [LLM] Error: {e}")
        return {}


def parse_fields(text: str) -> dict[str, str]:
    """Extract field values from preprocessed text. Handles 'Field: value', 'Field - value', table rows."""
    text = text or ""
    result = {f: "" for f in MANDATORY_FIELDS}
    lines = [ln.strip() for ln in text.replace("\r", "\n").split("\n") if ln.strip()]
    for i, line in enumerate(lines):
        for field in MANDATORY_FIELDS:
            if result[field]:
                continue
            if not line.lower().startswith(field.lower()):
                continue
            for sep in (":", "-"):
                if sep in line:
                    val = line.split(sep, 1)[1].strip()
                    max_len = 8000 if field == "Job Description" else 1000
                    result[field] = val[:max_len]
                    break
            if not result[field] and i + 1 < len(lines) and not _is_field_label(lines[i + 1]):
                result[field] = lines[i + 1][:1000]
            break
    for field in MANDATORY_FIELDS:
        if result[field]:
            continue
        escaped = re.escape(field)
        m = re.search(rf"{escaped}\s*[:\-]\s*(.+?)(?=\n\s*[A-Z]|\n\n|$)", text, re.DOTALL | re.I)
        if m:
            val = re.sub(r"\s+", " ", m.group(1).strip())
            max_len = 8000 if field == "Job Description" else 1000
            result[field] = val[:max_len]
    return result


def _count_empty(fields: dict) -> int:
    return sum(1 for v in fields.values() if not (v and str(v).strip()))


async def extract_and_save(input_dir: Path, output_path: Path, use_llm: bool = True) -> None:
    files = sorted(p for p in input_dir.iterdir() if p.suffix.lower() in WORD_EXTENSIONS)
    if not files:
        print(f"No .doc/.docx files in {input_dir}")
        return
    print(f"Extracting and preprocessing {len(files)} file(s)...")
    if use_llm and not os.getenv("CORE_API_KEY"):
        print("  (Set CORE_API_KEY in .env to use LLM for field extraction)")
    rows = []
    for i, file_path in enumerate(files):
        content = file_path.read_bytes()
        text = extract_text(content, file_path.name)
        fields_from_tables = _extract_fields_from_tables(content) if file_path.suffix.lower() == ".docx" else {}
        fields_from_text = parse_fields(text)
        fields = {}
        for f in MANDATORY_FIELDS:
            fields[f] = fields_from_tables.get(f) or fields_from_text.get(f) or ""
        do_llm = use_llm and bool(os.getenv("CORE_API_KEY"))
        if do_llm:
            print(f"  [{i+1}/{len(files)}] {file_path.name} - calling LLM...")
            llm_fields = await _extract_fields_with_llm(text, verbose=True)
            for f in MANDATORY_FIELDS:
                if not fields[f] and llm_fields.get(f):
                    fields[f] = llm_fields[f]
        else:
            print(f"  [{i+1}/{len(files)}] {file_path.name}")
        if _count_empty(fields) == len(MANDATORY_FIELDS) and text:
            fields["Job Description"] = text[:8000]
        row = {"fileName": file_path.name, "extractedText": text}
        row.update(fields)
        rows.append(row)
    df = pd.DataFrame(rows)
    key_cols = ["fileName", "Organization Name", "Job Code", "Position title", "Job Title (from Source Job Description)", "Job Description", "Base Salary", "Job Location"]
    rest = [c for c in MANDATORY_FIELDS if c not in key_cols[1:]]
    col_order = key_cols + rest + ["extractedText"]
    df = df[[c for c in col_order if c in df.columns]]
    df.to_excel(output_path, index=False, engine="openpyxl")
    print(f"Wrote {len(rows)} rows to {output_path}")


def main() -> None:
    p = argparse.ArgumentParser()
    p.add_argument("-i", "--input-dir", type=Path, required=True, help="Folder with Word docs")
    p.add_argument("-o", "--output", type=Path, default=Path("extracted_output.xlsx"), help="Output Excel")
    p.add_argument("--no-llm", action="store_true", help="Skip LLM, use table/text parsing only")
    args = p.parse_args()
    if not args.input_dir.exists():
        print(f"Directory not found: {args.input_dir}")
        exit(1)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    asyncio.run(extract_and_save(args.input_dir, args.output, use_llm=not args.no_llm))


if __name__ == "__main__":
    main()
